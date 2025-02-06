import tkinter as tk
from tkinter import ttk
import threading
import tkinter.messagebox
import docx
import time
from tkinter import filedialog
import os
import pyglet
import pyzipper
import shutil
from PIL import Image,ImageTk

def clear(x=0,y=0):
    cl="""
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                0
                                                                                                                
"""
    lc=(tk.Label(root,text=cl,font=('等线',50)))
    lc.place(x=x,y=y)

def docx_replace_text(document,old_word, new_word):
    
    for x in document.paragraphs:
        if old_word in x.text:  # t 尽量短，一个最好，不然这里可能会被拆分 如果替换失败 DEBUG这里查看x.text
            inline = x.runs  # t 修改runs中的字符串 可以保留格式
            for i in range(len(inline)):
                if old_word in inline[i].text:
                    text = inline[i].text.replace(old_word, new_word)
                    inline[i].text = text
                    #print("成功替换文本 正文",old_word,"->",new_word)

    
    children = document.element.body.iter()
    #文本框 图形
    for child in children:
        if child.tag.endswith('txbx'):
            for ci in child.iter():
                if ci.tag.endswith('main}t'):
                    if old_word in ci.text:
                        ci.text = new_word
                        #print("成功替换文本 文本框 图形",old_word,"->",new_word)
    #表格
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                # 遍历单元格中的所有段落
                for paragraph in cell.paragraphs:
                    # 遍历段落中的所有运行
                    for run in paragraph.runs:
                        if old_word in run.text:
                            run.text = run.text.replace(old_word, new_word)
    #页眉与页脚
    for section in document.sections:
        header=section.header
        footer=section.footer
        if header:
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    run.text=run.text.replace(old_word, new_word)
                    #print("成功替换页眉",old_word,"->",new_word)
        if footer:
            for paragraph in footer.paragraphs:
                for run in paragraph.runs:
                    run.text=run.text.replace(old_word, new_word)
                    #print("成功替换页脚",old_word,"->",new_word)
    return document

def ReadFile(FilePath,encoding="utf-8"):
    try:
        with open(FilePath,"r+",encoding=encoding) as file:
            return file.read()
    except UnicodeDecodeError:
        #print("解码失败",encoding)
        if encoding=="utf-8":
            encoding="gbk"
            return ReadFile(FilePath,encoding=encoding)
        elif encoding=="gbk":
            encoding="utf-16"
            return ReadFile(FilePath,encoding=encoding)
        elif encoding=="utf-16":
            encoding="latin-1"
            return ReadFile(FilePath,encoding=encoding)
        else:
            print(FilePath,"解码失败，没有找到合适的解码方式")

def extract_zip_with_password(zip_file, extract_to, password):
    try:
        with pyzipper.AESZipFile(zip_file) as zf:
            zf.setpassword(password.encode('utf-8'))
            zf.extractall(path=extract_to)
        return 1
    except RuntimeError:
        return -1


def is_zip_encrypted(zip_file):
    with pyzipper.AESZipFile(zip_file, 'r') as zf:
        for info in zf.infolist():
            if info.flag_bits & 0x1:
                return True
    return False

class GenerateStart_:
    def __init__(self):
        pass
    def do1(self):
        global NameOfInsuranceType,NameOfInstitutionType,InfoLabel
        
        
        if CustomerNameEntry.get()=="":
            tkinter.messagebox.showinfo("提示","请填写客户名称！")
            return None
        if InsuranceTypeCombobox.get()=="":
            tkinter.messagebox.showinfo("提示","选择险种！")
            return None
        if InsuranceTypeCombobox.get() not in NameOfInsuranceType:
            tkinter.messagebox.showinfo("提示","未知险种！")
            return None
        if InstitutionTypeCombobox.get()=="":
            tkinter.messagebox.showinfo("提示","选择机构！")
            return None
        if InstitutionTypeCombobox.get() not in NameOfInstitutionType:
            tkinter.messagebox.showinfo("提示","未知机构！")
            return None

        clear()
        InfoLabel=(tk.Label(root,text="正在准备生成",font=('等线',30)))
        InfoLabel.place(x=90,y=90)

        if os.path.exists("./tmp"):
            try:
                shutil.rmtree("./tmp")
            except PermissionError:
                InfoLabel["text"]="缓存目录异常！\n请手动删除 tmp 文件夹。"
                tkinter.messagebox.showerror("错误","缓存目录异常！\n请手动删除 tmp 文件夹。")
                return None
        os.mkdir("tmp")
        hide_folder("./tmp")

        print(CustomerNameEntry.get(),InsuranceTypeCombobox.get())

        if os.path.exists("./data.zip")==False:
            InfoLabel["text"]="没有找到压缩包！\n请检查压缩包存在后再次运行。"
            tkinter.messagebox.showerror("错误","没有找到压缩包！\n请检查压缩包存在后再次运行。")
            return None


        InfoLabel["text"]="正在解密文件"
        
        result=is_zip_encrypted("./data.zip")
        if result==False:
            InfoLabel["text"]="压缩包异常！\n没有检测到密码，\n请设置密码后再次运行。"
            tkinter.messagebox.showerror("错误","压缩包异常！\n没有检测到密码，请设置密码后再次运行。")
            return None
        
        result=extract_zip_with_password("./data.zip","./tmp","CCIC")
        if result==-1:
            InfoLabel["text"]="压缩包异常！\n密码错误，请检查解压密码后再次运行。"
            tkinter.messagebox.showerror("错误","压缩包异常！\n密码错误，请检查解压密码后再次运行。")
            return None

        InfoLabel["text"]="正在读取数据"

        if os.path.exists("./tmp/模板")==False:
            tkinter.messagebox.showerror("错误","文件目录错误！\n请重新检查压缩包！")
            return None
        if os.path.exists("./tmp/替换文本/保险责任")==False:
            tkinter.messagebox.showerror("错误","文件目录错误！\n请重新检查压缩包！")
            return None
        if os.path.exists("./tmp/替换文本/除外责任")==False:
            tkinter.messagebox.showerror("错误","文件目录错误！\n请重新检查压缩包！")
            return None
        
        if os.path.exists("./tmp/替换文本/保险责任/%s.txt"%(InsuranceTypeCombobox.get()))==False:
            tkinter.messagebox.showerror("错误","目标文件不存在！")
            return None
        if os.path.exists("./tmp/替换文本/除外责任/%s.txt"%(InsuranceTypeCombobox.get()))==False:
            tkinter.messagebox.showerror("错误","目标文件不存在！")
            return None
        if os.path.exists("./tmp/模板/模板.docx")==False:
            tkinter.messagebox.showerror("错误","模板文件不存在！")
            return None

        document = docx.Document("./tmp/模板/模板.docx")

        ReplaceRules=[("**客户名称**",CustomerNameEntry.get()),
                      ("**险种**",InsuranceTypeCombobox.get()),
                      ("**日期**",time.strftime('%Y年%m月%d日')),
                      ("**机构**",InstitutionTypeCombobox.get()),
                      ("**年**",time.strftime('%y年')),
                      ("**月**",time.strftime('%m月')),
                      ("**日**",time.strftime('%d日'))
                      ]
        
        with open("./tmp/替换文本/保险责任/%s.txt"%(InsuranceTypeCombobox.get()),"r+",encoding="utf-8") as file:
            ReplaceRules.append(["**保险责任**",file.read()])
        with open("./tmp/替换文本/除外责任/%s.txt"%(InsuranceTypeCombobox.get()),"r+",encoding="utf-8") as file:
            ReplaceRules.append(["**除外责任**",file.read()])
        with open("./tmp/替换文本/投保项目/%s.txt"%(InsuranceTypeCombobox.get()),"r+",encoding="utf-8") as file:
            ReplaceRules.append(["**投保项目**",file.read()])

        InfoLabel["text"]="已完成 0%"
        root.update_idletasks()

        for i in range(len(ReplaceRules)):
            document=docx_replace_text(document,ReplaceRules[i][0],ReplaceRules[i][1])
            InfoLabel["text"]="已完成 "+str(round(100.0/len(ReplaceRules)*i,2))+"%"
            root.update_idletasks()
        if os.path.exists("./tmp"):
            shutil.rmtree("./tmp")
            
        InfoLabel["text"]="已完成 100%"
        root.update_idletasks()
        
        FilePath=filedialog.asksaveasfilename(title=u'保存文件',filetypes=[("Word 文档","*.docx"),
                                                                       ('All Files', '*.*')],
                                              defaultextension="*.docx",
                                              initialfile=CustomerNameEntry.get()+" 保险方案.docx")

        if FilePath=="":
            clear()
            InfoLabel["text"]="请保存文件！"
            InfoLabel.place(x=90,y=90)
            time.sleep(0.2)
            tkinter.messagebox.showerror("错误","请保存文件！")
            FilePath=filedialog.asksaveasfilename(title=u'保存文件',filetypes=[("Word 文档","*.docx"),
                                                                           ('All Files', '*.*')],
                                                  defaultextension="*.docx",
                                                  initialfile=CustomerNameEntry.get()+" 保险方案.docx")
    
        if FilePath:
            try:
                document.save(FilePath)
            except FileNotFoundError:
                clear()
                InfoLabel["text"]="文件名或文件路径错误！"
                InfoLabel.place(x=90,y=90)
                time.sleep(0.2)
                tkinter.messagebox.showerror("错误","文件名或文件路径错误！")
            except PermissionError:
                clear()
                InfoLabel["text"]="文件被占用导致无法保存！\n建议关闭文件或更换文件名。"
                InfoLabel.place(x=90,y=90)
                time.sleep(0.2)
                tkinter.messagebox.showerror("错误","文件被占用导致无法保存\n建议关闭文件或更换文件名！")
            
        else:
            tkinter.messagebox.showerror("错误","未能成功保存文件。\n没有选择路径！")

        
        clear()
        main()
    def start(self):
        com = threading.Thread(target=self.do1)
        com.daemon = 1
        com.start()
        
        
def GenerateStart(_=None):
    GenerateStart_().start()
    
def hide_folder(folder_path):
    # 使用 attrib 命令设置隐藏属性
    os.system(f'attrib +h "{folder_path}"')

def cut_image(image,x=0,y=0,magnification=1):
    """全部参数 image,x,y,magnification
image为PIL格式"""
    if x==0 and y==0 and magnification==1:
        print("必须填入一个x值或y值或缩放倍率，且不为0")
    if x!=0 and y!=0:
        print("只能填入一个x值或y值或缩放倍率，且不为0")
    (width, height) = image.size
    if magnification!=1:
        image.thumbnail((int(width*magnification), int(height*magnification)))
    else:
        if x!=0:
            magnification=x/width
            image.thumbnail((int(width*magnification), int(height*magnification)))
        else:
            magnification=y/height
            image.thumbnail((int(width*magnification), int(height*magnification)))
    return image


def main():
    global CustomerNameEntry,InsuranceTypeCombobox,InstitutionTypeCombobox
    global NameOfInsuranceType,NameOfInstitutionType,InfoLabel
    global photo

    InfoLabel=(tk.Label(root,text="正在准备生成",font=('等线',30)))
    InfoLabel.place(x=90,y=90)
    
    if os.path.exists("./data.zip")==False:
        InfoLabel["text"]="没有找到压缩包！\n请检查压缩包存在后再次运行。"
        tkinter.messagebox.showerror("错误","没有找到压缩包！\n请检查压缩包存在后再次运行。")
        return None

    InfoLabel["text"]="正在解密文件"
        
    result=is_zip_encrypted("./data.zip")
    if result==False:
        InfoLabel["text"]="压缩包异常！\n没有检测到密码，\n请设置密码后再次运行。"
        tkinter.messagebox.showerror("错误","压缩包异常！\n没有检测到密码，请设置密码后再次运行。")
        return None
        
    result=extract_zip_with_password("./data.zip","./tmp","CCIC")
    if result==-1:
        InfoLabel["text"]="压缩包异常！\n密码错误，请检查解压密码后再次运行。"
        tkinter.messagebox.showerror("错误","压缩包异常！\n密码错误，请检查解压密码后再次运行。")
        return None

    try:
        image = Image.open(r'./image/1.jpg')
        image = cut_image(image,x=800)
        photo = ImageTk.PhotoImage(image)
        canvas = tk.Canvas(root, width=photo.width(), height=photo.height())
        canvas.create_image(0, 0, anchor=tk.NW, image=photo)
        canvas.create_text(50, 290, text="客户名称：", font=("等线", 17),anchor=tk.W)
        canvas.create_text(50, 340, text="险种：", font=("等线", 17),anchor=tk.W)
        canvas.create_text(50, 390, text="机构：", font=("等线", 17),anchor=tk.W)

        canvas.place(x=0,y=0)
    except FileNotFoundError:
        tkinter.messagebox.showwarning("警告", "图片文件未找到！")
    except Exception as e:
        tkinter.messagebox.showwarning("警告", f"图片显示异常！{e}")

    CustomerNameEntry=(tk.Entry(root,font=('等线',20),width=40))
    CustomerNameEntry.place(x=170,y=275)

    NameOfInsuranceType=[]
    NameOfInstitutionType=[]

    try:
        NameOfInsuranceType=ReadFile("./tmp/其它文件/险种名称.txt").split("\n")
    except FileNotFoundError:
        tkinter.messagebox.showerror("错误","险种名称 不存在！")
    try:
        NameOfInstitutionType=ReadFile("./tmp/其它文件/机构.txt").split("\n")
    except FileNotFoundError:
        tkinter.messagebox.showerror("错误","机构.txt 不存在！")

    if os.path.exists("./tmp"):
        shutil.rmtree("./tmp")
    InfoLabel.place_forget()
        
    InsuranceTypeCombobox = ttk.Combobox(root,values=NameOfInsuranceType,width=20,height=20,font=('等线',20))
    InsuranceTypeCombobox.place(x=170,y=325)

    InstitutionTypeCombobox = ttk.Combobox(root,values=NameOfInstitutionType,width=20,height=20,font=('等线',20))
    InstitutionTypeCombobox.place(x=170,y=375)

    b1=tk.Button(root,text='生成',font=('等线',15),width=15,height=3,command=GenerateStart)
    b1.place(x=500,y=325)

    label=(tk.Label(root,text="软件制作单位：嘉兴中心支公司",font=('等线',15)))
    label.place(x=450,y=450)



root = tk.Tk()
root.title('CCIC保险方案简易生成器 1.0')
root.geometry("800x550+50+50")
pyglet.options['win32_gdi_font'] = True
try:
    pyglet.font.add_file('./等线.ttf')
except FileNotFoundError:
    try:
        pyglet.font.add_file('等线.ttf')
    except FileNotFoundError:
        tkinter.messagebox.showwarning("警告","导入字体异常！\n显示可能错误！")

if os.path.isfile("logo.ico"):
    root.iconbitmap("logo.ico")

if os.path.exists("./tmp"):
    try:
        shutil.rmtree("./tmp")
    except PermissionError:
        tkinter.messagebox.showerror("错误","缓存目录异常！\n请手动删除 tmp 文件夹。")
os.mkdir("tmp")
hide_folder("./tmp")
    
root.bind("<Return>",GenerateStart)

main()

root.mainloop()
